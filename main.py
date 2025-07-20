from docx import Document

def gen_resume():
    # Get user input
    Name = input("Enter your Full Name: ")
    EmailId = input("Enter your Email Id: ")
    PhoneNo = input("Enter your Phone Number: ")
    LinkedIn = input("Enter your LinkedIn URL: ")
    About = input("Enter a short description About Yourself: ")

    # Skills and education
    Skills = input("List your skills (comma separated): ").split(",")
    Education = input("Enter your highest qualification: ")
    University = input("Enter your university/school name: ")

    # Experience
    Experience_Count = int(input("How many jobs do you want to add? "))
    Experience = []
    for i in range(Experience_Count):
        Job_Title = input(f"\nJob Title #{i+1}: ")
        Company = input("Company: ")
        Duration = input("Duration (e.g., Jan 2021 - Dec 2023): ")
        Desc = input("Job Description: ")
        Experience.append((Job_Title, company, duration, Desc))

    # Create document
    doc = Document()
    doc.add_heading(Name, 0)

    # Contact info
    doc.add_paragraph(f"{EmailId} | {PhoneNo} | {LinkedIn}")

    # Summary
    doc.add_heading("Summary", level=1)
    doc.add_paragraph(About)

    # Education
    doc.add_heading("Education", level=1)
    doc.add_paragraph(f"{Education} - {University}")

    # Experience
    doc.add_heading("Work Experience", level=1)
    for title, company, duration, desc in Experience:
        doc.add_paragraph(f"{title} at {company} ({duration})", style='List Bullet')
        doc.add_paragraph(desc)

    # Skills
    doc.add_heading("Skills", level=1)
    doc.add_paragraph(", ".join([skill.strip() for skill in Skills]))

    # Save
    filename = Name.replace(" ", "_").lower() + "_resume.docx"
    doc.save(filename)
    print(f"\nYour resume saved as {filename}")

# Run
gen_resume()
